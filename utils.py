# utils.py
# -*- coding: utf-8 -*-
import os
import json
import re
from pathlib import Path
import google.generativeai as genai
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from dotenv import load_dotenv # Kept for potential local testing, though unused in Streamlit UI
import pandas as pd
from docx import Document # Kept, but the function is not called in streamlit_app.py
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Configuration and Constants ---

EXPECTED_KEYS = [
    "main_subject", "topic_breakdown", "key_vocabulary",
    "formulas_and_principles", "teacher_insights",
    "exam_focus_points", "common_mistakes_explained"
]

# --- NEW PROMPT WITH HIGHLIGHTING INSTRUCTION ---
SYSTEM_PROMPT = """
You are a master academic analyst creating a concise, timestamped study guide from a video transcript file. The transcript text contains timestamps in formats like (MM:SS) or [HH:MM:SS].

**Primary Goal:** Create a detailed summary. For any key point you extract, you MUST find its closest preceding timestamp in the text and include it in your response as total seconds.

**Instructions:**
1.  Analyze the entire transcript.
2.  For every piece of information you extract, find the nearest timestamp that comes *before* it in the text. Convert that timestamp into **total seconds** (e.g., (01:30) becomes 90).
3.  **Highlighting:** Inside any 'detail', 'definition', 'explanation', or 'insight' string, find the single most critical phrase (3-5 words) and wrap it in `<hl>` and `</hl>` tags. For example: 'The derivative is a <hl>rate of change</hl>.' Do this only once per item where appropriate.
4.  Be concise. Each point must be a short, clear sentence.
5.  Extract the following information:
    -   main_subject: A short phrase identifying the main subject.
    -   topic_breakdown: For each topic, list details as objects with "detail" and "time" keys.
    -   key_vocabulary: A list of objects with "term", "definition", and "time" keys.
    -   formulas_and_principles: A list of objects with "formula_or_principle", "explanation", and "time" keys.
    -   teacher_insights: A list of objects with "insight" and "time" keys.
    -   exam_focus_points: A list of objects with "point", and "time" keys.
    -   common_mistakes_explained: A list of objects with "mistake", "explanation", and "time" keys.
6.  Return your output **only** as a single, valid JSON object.

The JSON structure must be exactly as follows:
{
  "main_subject": "Subject Name",
  "topic_breakdown": [{"topic": "Topic 1", "details": [{"detail": "This is a <hl>short detail</hl>.", "time": 120}]}],
  "key_vocabulary": [{"term": "Term 1", "definition": "A <hl>short definition</hl>.", "time": 150}],
  "formulas_and_principles": [{"formula_or_principle": "Principle 1", "explanation": "A <hl>brief explanation</hl>.", "time": 180}],
  "teacher_insights": [{"insight": "<hl>Short insight</hl> 1.", "time": 210}],
  "exam_focus_points": [{"point": "Brief <hl>focus point</hl> 1.", "time": 240}],
  "common_mistakes_explained": [{"mistake": "Mistake 1", "explanation": "A <hl>short explanation</hl>.", "time": 270}]
}
"""

# --- Color Palette (R, G, B) ---
COLORS = {
    "title_bg": (40, 54, 85), "title_text": (255, 255, 255),
    "heading_text": (40, 54, 85), "link_text": (0, 0, 255), # Blue
    "body_text": (30, 30, 30), "line": (220, 220, 220),
    "highlight_bg": (255, 255, 0) # Yellow
}

# --- Helper Functions ---

def get_video_id(url: str) -> str | None:
    patterns = [
        r"(?<=v=)[^&#?]+", r"(?<=be/)[^&#?]+", r"(?<=live/)[^&#?]+",
        r"(?<=embed/)[^&#?]+", r"(?<=shorts/)[^&#?]+"
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match: return match.group(0)
    return None

def clean_gemini_response(response_text: str) -> str:
    match = re.search(r'```json\s*(\{.*?\})\s*```|(\{.*?\})', response_text, re.DOTALL)
    if match: return match.group(1) if match.group(1) else match.group(2)
    return response_text.strip()

def summarize_with_gemini(api_key: str, transcript_text: str) -> dict | None:
    # Print statements are kept for debugging purposes in Streamlit logs
    print("    > Sending transcript to Gemini API...")
    try:
        genai.configure(api_key=api_key); model = genai.GenerativeModel('gemini-pro-latest')
        response = model.generate_content(f"{SYSTEM_PROMPT}\n\nTranscript:\n---\n{transcript_text}")
        cleaned_response = clean_gemini_response(response.text)
        return json.loads(cleaned_response)
    except Exception as e:
        print(f"    > An error occurred with the API call: {e}"); return None

def format_timestamp(seconds: int) -> str:
    minutes = seconds // 60; seconds = seconds % 60
    return f"[{minutes:02}:{seconds:02}]"

# --- Advanced PDF Class with Local Font Support ---

class PDF(FPDF):
    def __init__(self, font_path, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.font_name = "NotoSans"
        # Font paths are relative to the execution directory in Streamlit Cloud
        self.add_font(self.font_name, "", str(font_path / "NotoSans-Regular.ttf"))
        self.add_font(self.font_name, "B", str(font_path / "NotoSans-Bold.ttf"))

    def create_title(self, title):
        self.set_font(self.font_name, "B", 24); self.set_fill_color(*COLORS["title_bg"]); self.set_text_color(*COLORS["title_text"])
        self.cell(0, 20, title, align="C", fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT); self.ln(10)

    def create_section_heading(self, heading):
        self.set_font(self.font_name, "B", 16); self.set_text_color(*COLORS["heading_text"])
        self.cell(0, 10, heading, new_x=XPos.LMARGIN, new_y=YPos.NEXT); self.set_draw_color(*COLORS["line"])
        self.line(self.get_x(), self.get_y(), self.get_x() + 190, self.get_y()); self.ln(5)

    def write_highlighted_text(self, text, style=''):
        """Parses <hl> tags and applies highlighting."""
        self.set_font(self.font_name, style, 11)
        self.set_text_color(*COLORS["body_text"])
        
        parts = re.split(r'(<hl>.*?</hl>)', text)
        for part in parts:
            if part.startswith('<hl>'):
                highlight_text = part[4:-5] # Get text inside tags
                self.set_fill_color(*COLORS["highlight_bg"])
                self.set_font(self.font_name, 'B', 11) # Bold for highlighted text
                self.cell(self.get_string_width(highlight_text), 7, highlight_text, fill=True)
                self.set_font(self.font_name, style, 11) # Reset to regular
            else:
                self.set_fill_color(255, 255, 255) # No background
                self.write(7, part)
        self.ln()

# --- Save to Excel Function (XLSX) ---
def save_to_excel(data: dict, output_path: Path):
    print(f"    > Saving to Excel: {output_path.name}")
    try:
        flat_data = []
        for key, values in data.items():
            if key == "main_subject" or not values: continue
            
            section_name = key.replace('_', ' ').title()
            for item in values:
                row = {"Section": section_name}
                if isinstance(item, dict):
                    if 'details' in item and isinstance(item['details'], list):
                        row["Topic"] = item.get('topic', '')
                        for detail in item['details']:
                            detail_row = row.copy()
                            detail_row["Detail"] = detail.get('detail', '').replace('<hl>', '').replace('</hl>', '')
                            detail_row["Time (s)"] = detail.get('time', 0)
                            flat_data.append(detail_row)
                    else:
                        for sk, sv in item.items():
                            row[sk.replace('_', ' ').title()] = str(sv).replace('<hl>', '').replace('</hl>', '')
                        flat_data.append(row)
                else:
                    flat_data.append({"Section": section_name, "Item": str(item)})
        
        df = pd.DataFrame(flat_data)
        # Use str(output_path) for compatibility with to_excel in cloud environment
        df.to_excel(str(output_path), index=False, sheet_name="Summary") 
        print("    > Excel file saved successfully.")
    except Exception as e:
        print(f"    > Error saving to Excel: {e}")

# --- Save to Word Function (kept, but unused in Streamlit UI) ---
def add_hyperlink(paragraph, text, url):
    """Adds a clickable hyperlink to a Word document paragraph."""
    # ... (code for adding hyperlink to DOCX) ...
    # This function is not called in streamlit_app.py but must exist in utils.py
    # because it is called by save_to_word
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF') # Blue color
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single') # Underline
    
    rPr.append(c)
    rPr.append(u)
    r.append(rPr)
    r.append(OxmlElement('w:t'))
    r.find(qn('w:t')).text = text
    
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def save_to_word(data: dict, video_id: str, output_path: Path):
    print(f"    > Saving to Word: {output_path.name}")
    base_url = f"https.www.youtube.com/watch?v={video_id}"
    try:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Noto Sans'
        style.font.size = Pt(11)
        
        # Title
        title = doc.add_paragraph(data.get("main_subject", "Video Summary"))
        title.style = doc.styles['Title']
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for key, values in data.items():
            if key == "main_subject" or not values: continue
            
            doc.add_heading(key.replace('_', ' ').title(), level=1)
            
            for item in values:
                is_nested = any(isinstance(v, list) for v in item.values())
                if is_nested:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item.get('topic', '')).bold = True
                    for detail_item in item.get('details', []):
                        timestamp_sec = int(detail_item.get('time', 0))
                        link = f"{base_url}&t={timestamp_sec}s"
                        
                        p_detail = doc.add_paragraph(style='List Bullet 2')
                        p_detail.add_run(detail_item.get('detail', ''))
                        p_detail.add_run(' ')
                        add_hyperlink(p_detail, format_timestamp(timestamp_sec), link)
                else:
                    timestamp_sec = int(item.get('time', 0))
                    link = f"{base_url}&t={timestamp_sec}s"
                    
                    p = doc.add_paragraph(style='List Bullet')
                    for i, (sk, sv) in enumerate(item.items()):
                        if sk != 'time':
                            p.add_run(f"{sk.replace('_', ' ').title()}: ").bold = True
                            p.add_run(f"{sv}")
                            if i < len(item) - 2: p.add_run("\n") 
                    
                    p.add_run(' ')
                    add_hyperlink(p, format_timestamp(timestamp_sec), link)
            doc.add_paragraph()
        
        doc.save(output_path)
        print("    > Word file saved successfully.")
    except Exception as e:
        print(f"    > Error saving to Word: {e}")

# --- Save to PDF Function (Primary Output) ---
def save_to_pdf(data: dict, video_id: str, font_path: Path, output_path: Path):
    print(f"    > Saving elegantly hyperlinked PDF: {output_path.name}")
    base_url = f"https.www.youtube.com/watch?v={video_id}"
    try:
        pdf = PDF(font_path=font_path)
        pdf.add_page()
        pdf.create_title(data.get("main_subject", "Video Summary"))
        font_name = pdf.font_name

        for key, values in data.items():
            if key == "main_subject" or not values: continue
            pdf.create_section_heading(key.replace('_', ' ').title())

            for item in values:
                is_nested = any(isinstance(v, list) for v in item.values())
                if is_nested:
                    pdf.set_font(font_name, "B", 11); pdf.set_text_color(*COLORS["body_text"])
                    pdf.multi_cell(0, 7, text=f"  {item.get('topic', '')}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    for detail_item in item.get('details', []):
                        timestamp_sec = int(detail_item.get('time', 0))
                        link = f"{base_url}&t={timestamp_sec}s"
                        display_text = f"    • {detail_item.get('detail', '')}"
                        
                        pdf.set_font(font_name, "", 11)
                        pdf.write_highlighted_text(display_text) # Use new highlight function
                        
                        pdf.set_text_color(*COLORS["link_text"]); pdf.cell(0, 7, text=format_timestamp(timestamp_sec), link=link, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="R")
                else:
                    timestamp_sec = int(item.get('time', 0)); link = f"{base_url}&t={timestamp_sec}s"
                    
                    pdf.set_font(font_name, "", 11)
                    for sk, sv in item.items():
                        if sk != 'time':
                            pdf.set_text_color(*COLORS["body_text"])
                            pdf.set_font(font_name, "B", 11)
                            pdf.write(7, f"• {sk.replace('_', ' ').title()}: ")
                            pdf.set_font(font_name, "", 11)
                            pdf.write_highlighted_text(str(sv))
                    
                    pdf.set_text_color(*COLORS["link_text"]); pdf.cell(0, 7, text=format_timestamp(timestamp_sec), link=link, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="R")
                pdf.ln(4)
            pdf.ln(5)

        # Use str() for output_path for compatibility
        pdf.output(str(output_path))
        print("    > PDF file saved successfully.")
    except Exception as e:
        print(f"    > Error saving to PDF: {e}")

# The file ends here. No main() or if __name__ == "__main__": block.
